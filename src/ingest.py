"""Ingestão de documentos do usuário no A.P.O.L.O.

Você solta um arquivo (texto, código, markdown ou PDF) e ele é fatiado em trechos,
indexado no ChromaDB (recall semântico) e registrado no Supabase. A partir daí o
A.P.O.L.O. responde sobre os SEUS documentos e os cita — mesma memória usada no chat.
"""

import hashlib
import io
import logging
import re

from src.topics import classify_sector

logger = logging.getLogger(__name__)

CHUNK_CHARS = 1500      # tamanho de cada trecho indexado
CHUNK_OVERLAP = 150     # sobreposição p/ não cortar ideias no meio
MAX_CHUNKS = 40         # teto por arquivo (evita inflar a base com um doc gigante)
MAX_TEXT_CHARS = 200_000


def extract_pdf_text(data: bytes) -> str:
    """Extrai texto de um PDF. Requer `pypdf` (import preguiçoso)."""
    from pypdf import PdfReader  # noqa: PLC0415 — opcional, só carrega se houver PDF
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_docx_text(data: bytes) -> str:
    """Extrai texto de um .docx. Requer `python-docx` (import preguiçoso)."""
    from docx import Document  # noqa: PLC0415 — opcional, só carrega se houver DOCX
    doc = Document(io.BytesIO(data))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # Tabelas também
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def chunk_text(text: str, size: int = CHUNK_CHARS, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = re.sub(r"\n{3,}", "\n\n", (text or "")).strip()
    chunks: list[str] = []
    i = 0
    step = max(1, size - overlap)
    while i < len(text) and len(chunks) < MAX_CHUNKS:
        piece = text[i:i + size].strip()
        if piece:
            chunks.append(piece)
        i += step
    return chunks


class DocumentIngestor:
    """Fatiamento + indexação de documentos do usuário."""

    def __init__(self, rag=None, knowledge_db=None):
        self.rag = rag
        self.knowledge_db = knowledge_db

    def ingest_text(self, filename: str, text: str, source: str = "") -> dict:
        text = (text or "").strip()[:MAX_TEXT_CHARS]
        if len(text) < 20:
            return {"ok": False, "error": "Conteúdo vazio ou muito curto para indexar."}

        filename = (filename or "documento").strip()[:120]
        chunks = chunk_text(text)
        if not chunks:
            return {"ok": False, "error": "Não consegui extrair conteúdo do documento."}

        sector = classify_sector(f"{filename} {text[:300]}")
        base = hashlib.md5((source or filename).encode("utf-8", "ignore")).hexdigest()[:8]
        # Fonte citável: a URL real (ingestão de link) ou um pseudo-URL de arquivo.
        src = source or f"arquivo://{filename}"

        # ChromaDB — um trecho por chunk, para recall semântico granular.
        indexed = 0
        if self.rag:
            for k, chunk in enumerate(chunks):
                try:
                    self.rag.add_example(
                        f"# {filename} (parte {k + 1}/{len(chunks)})\n"
                        f"Fonte: {src}\n\n{chunk}",
                        f"userdoc_{base}_{k}",
                    )
                    indexed += 1
                except Exception as e:
                    logger.warning(f"[ingest] ChromaDB chunk {k}: {e}")

        # Supabase — um registro com a prévia do documento (não infla com N linhas).
        if self.knowledge_db:
            try:
                self.knowledge_db.save(
                    f"📄 {filename}", source or f"file://{filename}", text[:8000], "user_doc", [sector],
                )
            except Exception as e:
                logger.warning(f"[ingest] Supabase: {e}")

        logger.info(f"[ingest] '{filename}' → {indexed} trechos (setor: {sector})")
        return {
            "ok": True,
            "filename": filename,
            "chunks": indexed,
            "chars": len(text),
            "sector": sector,
        }
